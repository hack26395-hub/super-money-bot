import os
import uuid
import logging
import tempfile
import time
import telebot
from telebot import types
import google.generativeai as genai
from pdf2docx import Converter
from docx import Document
from docx.shared import RGBColor, Italic
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from deep_translator import GoogleTranslator
import traceback

# ========== التوكنات والمفاتيح ==========
TELEGRAM_TOKEN = "8433118363:AAH0iqeZVo3xz-KP_KQ7LxHSdhRZnOmb2LQ"
GEMINI_API_KEY = "AIzaSyCChd6IL-8hi9ttKOIwH-vVF57MzK8X26s"

# ========== تهيئة البوت و Gemini ==========
bot = telebot.TeleBot(TELEGRAM_TOKEN, parse_mode=None)
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# ========== إعدادات الترجمة ==========
translator = GoogleTranslator(source='en', target='ar')

# ========== مجلد مؤقت ==========
TEMP_DIR = tempfile.gettempdir()
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ========== دوال مساعدة ==========
def random_filename(ext=""):
    """توليد اسم ملف عشوائي باستخدام UUID"""
    return os.path.join(TEMP_DIR, uuid.uuid4().hex + ext)

def translate_text(text):
    """ترجمة نص من الإنجليزية إلى العربية"""
    if not text or not text.strip():
        return ""
    try:
        return translator.translate(text)
    except Exception as e:
        logger.error(f"خطأ في الترجمة: {e}")
        return "[ترجمة غير متاحة حالياً]"

def set_paragraph_style(paragraph, color_rgb=(31, 73, 125), italic=True):
    """تطبيق اللون الأزرق والخط المائل على فقرة"""
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
        run.font.italic = italic
    if not paragraph.runs:
        run = paragraph.add_run()
        run.font.color.rgb = RGBColor(*color_rgb)
        run.font.italic = italic
    return paragraph

def add_arabic_translation_below_paragraph(original_para, doc):
    """إضافة فقرة ترجمة أسفل الفقرة الأصلية"""
    eng_text = original_para.text.strip()
    if not eng_text:
        return
    arabic_text = translate_text(eng_text)
    # إضافة فقرة جديدة أسفل الفقرة الأصلية مباشرة
    new_para = doc.add_paragraph()
    new_para.add_run(arabic_text)
    set_paragraph_style(new_para)
    return new_para

def translate_docx_file(docx_path, output_path):
    """
    يقوم بترجمة جميع الفقرات في ملف DOCX وإضافة الترجمة أسفل كل فقرة إنجليزية.
    يحافظ على الصور والجداول (لا يعدلها، فقط يتعامل مع الفقرات).
    """
    doc = Document(docx_path)
    # قائمة لتخزين الفقرات التي تمت معالجتها (نمر على الفقرات الأصلية)
    paragraphs = list(doc.paragraphs)
    # نضيف الترجمة أسفل كل فقرة (هذا يزيد عدد الفقرات أثناء التكرار، لذلك نتعامل بحذر)
    # سنقوم بتجميع الفقرات الجديدة في نهاية المستند لأن الإضافة المباشرة تغير الفهارس.
    # لكن المتطلب "أسفل كل سطر" يمكن تحقيقه بإضافة فقرة بعد كل فقرة موجودة.
    # طريقة آمنة: نعمل على قائمة مؤقتة ثم نعيد ترتيبها؟ أسهل: المرور على الفقرات في كل Loop
    # ولكن إضافة فقرة جديدة داخل الحلقة ستجعلها تظهر أسفل الفقرة الحالية (add_paragraph يضيف في النهاية).
    # لكي تظهر أسفل الفقرة مباشرة، نحتاج إلى إدراجها بعد الفقرة الحالية.
    # سنقوم بإدراج فقرة جديدة في موضع (index+1) باستخدام _element.
    counter = 0
    for para in paragraphs:
        # نتخطى الفقرات الفارغة أو التي تحتوي على مسافات فقط
        if para.text.strip():
            # إدراج فقرة جديدة بعد الفقرة الحالية
            new_para = doc.add_paragraph()
            # نقل محتوى الفقرة الجديدة إلى الموضع الصحيح (بعد الفقرة الأصلية)
            # الطريقة: تحديد العنصر الأصل وإدراج العنصر الجديد بعده
            parent = para._element.getparent()
            index = parent.index(para._element)
            new_para_element = new_para._element
            parent.insert(index + 1, new_para_element)
            # الآن نضع الترجمة في هذه الفقرة
            eng_text = para.text.strip()
            arabic_text = translate_text(eng_text)
            new_para.clear()
            run = new_para.add_run(arabic_text)
            run.font.color.rgb = RGBColor(31, 73, 125)
            run.font.italic = True
    # حفظ الملف الجديد
    doc.save(output_path)
    return output_path

def convert_pdf_to_docx(pdf_path, docx_path):
    """تحويل PDF إلى DOCX مع الحفاظ على الصور والجداول باستخدام pdf2docx"""
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    return docx_path

def process_document_file(file_path, original_extension, chat_id, message_id):
    """
    معالجة الملف: إذا كان PDF يحول إلى DOCX ثم يترجم، وإذا كان DOCX يترجم مباشرة.
    تعيد مسار الملف الناتج
    """
    try:
        # إرسال رسالة انتظار مؤقتة
        wait_msg = bot.send_message(chat_id, "⏳ جاري معالجة الملف... Please wait...", reply_to_message_id=message_id)

        unique_name = random_filename("")
        if original_extension == ".pdf":
            temp_docx = unique_name + ".docx"
            logger.info(f"تحويل PDF إلى DOCX: {file_path} -> {temp_docx}")
            convert_pdf_to_docx(file_path, temp_docx)
            output_docx = random_filename("_translated.docx")
            translate_docx_file(temp_docx, output_docx)
            os.remove(temp_docx)  # حذف الملف المؤقت بعد الترجمة
        elif original_extension == ".docx":
            output_docx = unique_name + "_translated.docx"
            translate_docx_file(file_path, output_docx)
        else:
            bot.edit_message_text("❌ نوع الملف غير مدعوم. أرسل PDF أو DOCX فقط.", chat_id, wait_msg.message_id)
            return None

        # حذف الملف الأصلي إذا كان تم تنزيله
        if os.path.exists(file_path):
            os.remove(file_path)

        # تعديل رسالة الانتظار إلى Done!
        bot.edit_message_text("✅ Done! تمت المعالجة بنجاح. جاري إرسال الملف...", chat_id, wait_msg.message_id)
        return output_docx
    except Exception as e:
        logger.error(f"خطأ في معالجة الملف: {traceback.format_exc()}")
        bot.send_message(chat_id, f"⚠️ حدث خطأ أثناء المعالجة: {str(e)}", reply_to_message_id=message_id)
        return None

# ========== معالجة الأوامر والرسائل ==========
@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    """رسالة الترحيب"""
    welcome_text = (
        "🎓 *أهلاً بك! أنا ذكاء قيس (Qais AI)*\n\n"
        "تم تطويري بواسطة المبرمج العبقري *قيس*، وأفخر بذلك!\n\n"
        "✨ *ماذا يمكنني أن أفعل لك؟*\n"
        "• أجب على أسئلتك العلمية (فيزياء، رياضيات، كيمياء) بأسلوب شرح مفصل.\n"
        "• أحوّل ملفات PDF إلى DOCX مع الحفاظ على الصور والجداول.\n"
        "• أترجم ملفات DOCX أو PDF (بعد التحويل) من الإنجليزية إلى العربية، وأضيف الترجمة أسفل كل سطر بلون أزرق وخط مائل.\n\n"
        "📤 *أرسل لي ملف PDF أو DOCX لترجمته، أو اسألني أي سؤال علمي.*"
    )
    bot.reply_to(message, welcome_text, parse_mode='Markdown')

@bot.message_handler(content_types=['document'])
def handle_document(message):
    """معالجة الملفات المرسلة (PDF, DOCX)"""
    try:
        file_info = bot.get_file(message.document.file_id)
        file_extension = os.path.splitext(message.document.file_name)[1].lower()
        if file_extension not in ['.pdf', '.docx']:
            bot.reply_to(message, "❌ أرسل ملف PDF أو DOCX فقط.")
            return

        downloaded_file = bot.download_file(file_info.file_path)
        unique_input = random_filename(file_extension)
        with open(unique_input, 'wb') as f:
            f.write(downloaded_file)

        # معالجة الملف
        output_file = process_document_file(unique_input, file_extension, message.chat.id, message.message_id)
        if output_file and os.path.exists(output_file):
            with open(output_file, 'rb') as f:
                bot.send_document(message.chat.id, f, caption="📄 *الملف بعد الترجمة* (تمت الإضافة أسفل كل سطر إنجليزي باللون الأزرق المائل)\n\n_أنا ذكاء قيس، فخور بخدمتك!_", parse_mode='Markdown')
            os.remove(output_file)
    except Exception as e:
        logger.error(f"خطأ في معالجة المستند: {e}")
        bot.reply_to(message, "⚠️ حدث خطأ غير متوقع. حاول مرة أخرى.")

@bot.message_handler(func=lambda message: True)
def handle_text(message):
    """الرد على الأسئلة النصية باستخدام Gemini (مع التركيز على العلوم)"""
    user_question = message.text.strip()
    if not user_question:
        return

    # إرسال مؤقت "يجري التفكير"
    wait_msg = bot.reply_to(message, "🧠 *ذكاء قيس يفكر...* يرجى الانتظار.", parse_mode='Markdown')

    try:
        # تحديد المجال العلمي؟ نطلب من Gemini الإجابة بشكل علمي مفصل
        prompt = f"""
        أنت "ذكاء قيس" (Qais AI)، تم تطويرك بواسطة المبرمج قيس. تتحدث بلسانه وتفخر بذلك.
        السؤال: {user_question}
        
        إذا كان السؤال متعلقاً بالفيزياء، الرياضيات، أو الكيمياء، قدم شرحاً مفصلاً وعلمياً.
        إذا كان السؤال غير علمي، قدم رداً عاماً ولكن بأسلوب مهذب ومفيد.
        قم بالتوقيع في نهاية الرد بـ "— ذكاء قيس (بتقنية Gemini 1.5 Flash)"
        """

        response = model.generate_content(prompt)
        answer = response.text if response.text else "عذراً، لم أستطع الإجابة على سؤالك حالياً."

        # تعديل رسالة الانتظار
        bot.edit_message_text(answer, chat_id=message.chat.id, message_id=wait_msg.message_id, parse_mode='Markdown')
    except Exception as e:
        logger.error(f"خطأ في استدعاء Gemini: {e}")
        bot.edit_message_text("⚠️ حدث خطأ تقني. يرجى المحاولة لاحقاً.", chat_id=message.chat.id, message_id=wait_msg.message_id)

# ========== تشغيل البوت بشكل مستمر (polling مع while True) ==========
def start_bot():
    """تشغيل البوت مع إعادة المحاولة التلقائية"""
    while True:
        try:
            logger.info("🔌 بدء تشغيل بوت ذكاء قيس...")
            bot.polling(none_stop=True, interval=1, timeout=60)
        except Exception as e:
            logger.error(f"⚠️ توقف البوت بسبب خطأ: {e}. إعادة التشغيل خلال 5 ثوانٍ...")
            time.sleep(5)

if __name__ == "__main__":
    start_bot()
